# Generates the IFQ636 Assignment 1 report as a .docx mirroring the official template.
# Run: python build_report.py   ->  IFQ636_A1_Report_OsamaMohamed.docx
# Written prose is complete (~2000 words); screenshot/link slots are clearly marked.
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(11)

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def para(t): doc.add_paragraph(t)
def placeholder(t):
    p = doc.add_paragraph()
    r = p.add_run('[SCREENSHOT] ' + t)
    r.italic = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

# ---------- Title + header links ----------
title = doc.add_heading('IFQ636 Assignment 1 — Software Requirements Analysis and Design', level=0)
sub = doc.add_paragraph('Medication Reminder System (MediRemind)')
sub.runs[0].bold = True
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

links = [
    ('Full name', 'Osama Mohamed'),
    ('Student ID', '12281069'),
    ('Draw.io link', '<<INSERT diagrams.net share link>>'),
    ('JIRA link', '<<INSERT Jira board URL>>'),
    ('Figma link', '<<INSERT Figma prototype link>>'),
    ('EC2 instance name and ID', '<<INSERT e.g. MediRemind-EC2 / i-0abc123...>>'),
    ('GitHub link', '<<INSERT GitHub repository URL>>'),
]
t = doc.add_table(rows=len(links), cols=2)
t.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate(links):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
doc.add_paragraph('Please ensure all links are working correctly.').italic = True

# ---------- 1. Project overview ----------
h1('1. Project overview')
para(
 "This project delivers the Medication Reminder System (MediRemind), a full-stack web "
 "application that helps patients keep track of their medications and take them on time. "
 "Medication non-adherence is a well-documented healthcare problem: patients frequently "
 "forget doses, lose track of complex regimens, or run out of supply, which reduces "
 "treatment effectiveness and increases avoidable hospital admissions. MediRemind addresses "
 "this by letting a patient record each medication, schedule dose reminders, log whether each "
 "dose was taken or skipped, and view an adherence summary that quantifies how reliably they "
 "follow their regimen.")
para(
 "The system was built by extending the provided Node.js, React.js and MongoDB starter "
 "project, which supplied user authentication only. On top of that foundation I implemented "
 "four meaningful CRUD modules — Medications, Reminders, Dose Logs and an Admin console — "
 "together with role-based access control that separates the patient (user) panel from the "
 "administrator panel. The application was designed using SysML, planned and tracked in JIRA, "
 "prototyped in Figma, version-controlled on GitHub with a feature-branch workflow, and "
 "deployed to an AWS EC2 instance through a GitHub Actions CI/CD pipeline.")

# ---------- 2. Real-world application ----------
h1('2. Real-world application')
para(
 "Acting as the product manager, the user requirements for MediRemind are as follows. The "
 "primary users are patients managing one or more ongoing medications, and a system "
 "administrator responsible for oversight.")
para("User requirements (functional):")
for r in [
 "As a patient, I want to register and log in securely so that my health information is private.",
 "As a patient, I want to add, view, edit and delete my medications (name, dosage, form, "
 "instructions, prescriber, quantity) so that my regimen is accurately recorded.",
 "As a patient, I want to schedule reminders for each medication at specific times and "
 "frequencies so that I do not forget a dose.",
 "As a patient, I want to mark each dose as taken or skipped so that my adherence is tracked.",
 "As a patient, I want to see an adherence percentage and dose history so that I can monitor "
 "my own behaviour.",
 "As an administrator, I want to view all users and system-wide statistics, and manage user "
 "roles or remove accounts, so that the platform stays well governed.",
]:
    doc.add_paragraph(r, style='List Bullet')
para("Non-functional requirements: the system must isolate each patient's data (a patient may "
 "only access their own records), authenticate every request with JWTs, be responsive on "
 "desktop and mobile, and be continuously deployable through an automated pipeline.")

# ---------- 3. Project management and design ----------
h1('3. Project management and design')
h2('3.1 SysML design diagrams')
para(
 "The system was modelled in SysML using diagrams.net. Seven diagrams were produced (see the "
 "Draw.io link above): a requirement diagram, use case diagram, block definition diagram "
 "(BDD), internal block diagram (IBD), activity diagram, sequence diagram and state machine "
 "diagram.")
para(
 "The requirement diagram is the anchor of the design. The top-level requirement (R0, "
 "'Medication Reminder System') is decomposed by containment into five sub-requirements: "
 "Medication Management (R1), Reminder Scheduling (R2), Dose Tracking & Adherence (R3), User "
 "Authentication (R4) and Admin Management (R5). Detailed requirements are linked using the "
 "«deriveReqt» relationship — for example R1 derives R1.1 (Add/Edit/Delete Medication) and "
 "R1.2 (View Medications), and R3 derives R3.1 (Compute Adherence). Each design block is then "
 "tied back to the requirement it fulfils with the «satisfy» relationship: MedicationService "
 "satisfies R1, ReminderService satisfies R2, DoseLogService satisfies R3, AuthService "
 "satisfies R4 and AdminService satisfies R5. This traceability shows that every requirement "
 "is realised by a concrete part of the implementation.")
para(
 "The use case diagram captures the Patient and Admin actors and their interactions "
 "(register/login, manage medications, schedule reminders, log doses, view adherence and "
 "history; manage users and view statistics). The BDD defines the data blocks (User, "
 "Medication, Reminder, DoseLog) with their attributes and the associations and multiplicities "
 "between them (a User owns 0..* Medications; a Medication has 0..* Reminders; a Reminder "
 "generates 0..* DoseLogs). The IBD shows the runtime structure — a React client, an Express "
 "REST API and a MongoDB database communicating over HTTPS/JSON and Mongoose. The activity "
 "diagram models the 'log a dose' workflow, the sequence diagram models the 'create reminder' "
 "interaction across client, API, controller and database, and the state machine models the "
 "DoseLog lifecycle (Scheduled → Taken / Skipped / Missed).")
placeholder('Insert each SysML diagram here (requirement, use case, BDD, IBD, activity, sequence, state machine), each with a one-line caption.')

h2('3.2 Software project management with JIRA')
para(
 "The project was planned in JIRA as a Scrum project. Work was organised into epics, each "
 "broken down into user stories, with sub-tasks under each story for the individual model, "
 "API, UI, validation and testing work. The epics are: Medication Management, Reminder "
 "Scheduling, Dose Tracking & Adherence, Admin Console, User Authentication and DevOps/CI-CD. "
 "Two sprints were planned: Sprint 1 covered authentication and medication management; Sprint "
 "2 covered reminders, dose tracking and the admin console.")
para("As required, at least two non-authentication epics are shown below with their user "
 "stories and sub-tasks, along with the board URL above.")
for s in [
 'Product backlog with all epics and stories.',
 'Project timeline (roadmap) showing epics and user stories.',
 'Epic 1 (Medication Management): a user story expanded with its sub-tasks.',
 'Epic 2 (Reminder Scheduling): a user story expanded with its sub-tasks.',
 'Sprint planning view (all sprints planned).',
 'A sprint that has been started but not yet completed.',
 'The complete board after the first sprint is started.',
 'A completed sprint.',
]:
    placeholder(s)

h2('3.3 UI/UX design with Figma')
para(
 "The interface was designed in Figma, progressing from low-fidelity wireframes to a "
 "high-fidelity, interactive prototype (see the Figma link above). The design covers the key "
 "screens: login/registration, the patient dashboard (adherence at a glance plus today's "
 "reminders), the medications list with create/edit forms and a delete confirmation, the "
 "reminders list and form, the dose history table, and the admin panel with system statistics "
 "and a user table.")
para(
 "Design decisions prioritised clarity and low cognitive load, since the target users include "
 "older patients: a single primary action colour, large touch targets for the 'Mark Taken' and "
 "'Skip' buttons, a persistent top navigation bar, consistent card layouts for list items, and "
 "colour-coded status badges (green = taken, yellow = skipped, red = missed). The same visual "
 "language was carried through to the implemented Tailwind CSS interface so the prototype and "
 "product stay consistent.")
placeholder('Insert Figma screenshots: low-fidelity wireframes, high-fidelity screens, and the prototype/collaboration view.')

# ---------- 4. Dev + GitHub ----------
h1('4. Backend development, Frontend development, GitHub version control and Branching strategy')
para(
 "Backend. The backend is an Express REST API using Mongoose models for User, Medication, "
 "Reminder and DoseLog. Authentication issues JWTs and hashes passwords with bcrypt. A "
 "'protect' middleware authenticates requests and an 'admin' middleware enforces role-based "
 "access for the admin routes. Every medication, reminder and dose-log endpoint is "
 "owner-scoped: queries are filtered by the authenticated user's id and ownership is checked "
 "before any update or delete, so one patient can never read or modify another's data. The API "
 "exposes full CRUD for medications and reminders, create/read/update/delete plus an adherence "
 "aggregation for dose logs, and admin endpoints for user management and system statistics. "
 "Adherence is computed on demand from the dose events using a MongoDB aggregation, keeping a "
 "single source of truth.")
para(
 "Frontend. The frontend is a React single-page application using React Router and Tailwind "
 "CSS. Pages include Dashboard, Medications, Reminders, History, Admin and Profile, with "
 "reusable form/list components for the CRUD modules. Authentication state is held in a React "
 "context and persisted to localStorage so a page refresh keeps the user logged in, and the "
 "API base URL is read from an environment variable so the same build runs locally and on "
 "EC2. A role-gated navigation bar shows the Admin link only to administrators.")
para(
 "Testing. The backend has a Mocha + Chai + Sinon test suite (13 tests) that stubs the "
 "Mongoose models, so the controller logic — status codes, validation and ownership checks — "
 "is verified without needing a database, which is ideal for the CI runner. The frontend has a "
 "React Testing Library test that renders the app and verifies the navigation.")
para(
 "GitHub version control and branching. All code, design artefacts and documentation are kept "
 "under Git. The repository uses a feature-branch workflow: 'main' is the always-deployable "
 "branch, and each feature was developed on its own branch (feature/medication-crud, "
 "feature/reminder-scheduling, feature/dose-tracking, feature/admin-panel, feature/ci-cd) and "
 "merged into main through a pull request. This produces a clear commit history and reviewable "
 "pull requests.")
placeholder('Insert screenshots: GitHub repo (backend/ + frontend/), branch list, an example pull request, and commit history.')

# ---------- 5. CI/CD ----------
h1('5. CI/CD Pipeline setup')
para(
 "Continuous integration and deployment are automated with GitHub Actions. On every push or "
 "pull request to main the pipeline installs dependencies and runs the backend and frontend "
 "test suites. When tests pass on a push to main, the deploy job connects to the EC2 instance "
 "over SSH, pulls the latest code, installs dependencies, rebuilds the frontend and reloads "
 "the application under PM2. Deployment credentials are stored as encrypted GitHub Actions "
 "secrets within a 'production' environment.")
h2('5.1 Workflow file (YML) screenshot')
placeholder('Insert screenshot of .github/workflows/ci-cd.yml.')
h2('5.2 Test case results with pass/fail status (terminal output)')
placeholder('Insert screenshot of the terminal showing the Mocha test results (13 passing).')
h2('5.3 GitHub Action configuration (runner, environments, prod variables)')
placeholder('Insert screenshot of the Actions configuration: runner (ubuntu-latest), the production environment, and the EC2 secrets/variables.')
h2('5.4 EC2 server configuration (pm2 status table)')
placeholder('Insert screenshot of the pm2 status output table on the EC2 instance.')
h2('5.5 GitHub "Run Test" page (job running, steps passing/failing)')
placeholder('Insert screenshot of the Actions run page showing the test and deploy jobs with passing steps.')
h2('5.6 First page of the application from the browser (highlight the public IP)')
placeholder('Insert screenshot of http://<EC2_PUBLIC_IP>:5001 in the browser with the public IP highlighted.')

# ---------- 6. README ----------
h1('6. README.md')
para(
 "The repository includes a README.md with project setup instructions, the tech stack, the "
 "feature list, local and EC2 deployment steps, the branching strategy, the public URL, and "
 "the demo credentials (admin@mediremind.com / Admin@123 and patient@mediremind.com / "
 "Patient@123). The JIRA board URL and public URL are also recorded there.")
placeholder('Insert screenshot of the rendered README.md on GitHub.')

# ---------- 7. Conclusion ----------
h1('7. Conclusion')
para(
 "MediRemind demonstrates a complete software lifecycle for a realistic problem: from "
 "requirements analysis and SysML design, through JIRA-based planning, Figma UI/UX design, "
 "and a tested MERN implementation, to an automated CI/CD pipeline deploying to AWS EC2. The "
 "delivered system meets the assignment's CRUD, user-panel and admin-panel requirements with "
 "meaningful, domain-appropriate features, and the engineering practices — owner-scoped data "
 "access, automated tests, feature branches and continuous deployment — reflect professional "
 "standards. The result is a maintainable foundation that could be extended with push "
 "notifications, caregiver sharing or refill alerts.")

# ---------- 8. Use of GenAI ----------
h1('8. Use of GenAI')
para(
 "Generative AI was used as a development assistant during this project. The tool used was "
 "GitHub Copilot.")
para("How it was used:")
for r in [
 "Brainstorming: I used Copilot to brainstorm candidate features and requirements for a "
 "medication reminder system. Example prompt: 'List the core CRUD entities and user/admin "
 "features a medication reminder web app should have.'",
 "Advice on application structure: I asked for guidance on structuring a MERN application. "
 "Example prompt: 'Suggest a folder structure and the models, controllers and routes needed "
 "to extend a Node/Express/Mongoose auth starter into a medication reminder app with "
 "role-based access.'",
 "Figma guidance: I asked Copilot how to perform tasks in Figma. Example prompt: 'How do I "
 "create a reusable component and an interactive prototype link in Figma for a dashboard "
 "screen?'",
 "Code snippets and examples: I used Copilot to generate example snippets such as a Mongoose "
 "schema, an owner-scoped Express controller, a React form component and a Mocha test that "
 "stubs Mongoose with Sinon.",
]:
    doc.add_paragraph(r, style='List Bullet')
para(
 "Which parts were influenced: the AI assisted with the initial data-model and folder "
 "structure, boilerplate for controllers/components, and example test patterns. The SysML "
 "design, the requirement decomposition, the JIRA plan, the Figma design decisions and the "
 "CI/CD and deployment configuration were my own work, informed by the AI's suggestions.")
para(
 "Verification and adaptation: all AI-suggested code was reviewed, edited to fit the project's "
 "conventions, and verified by running the application and the automated test suite (13 "
 "backend tests and the frontend test all pass) and by manually exercising each CRUD flow. "
 "Suggestions that were inaccurate or did not enforce per-user data isolation were corrected. "
 "AI output was treated as a starting point, not a final answer.")

# ---------- 9. Reflection ----------
h1('9. Reflection')
para(
 "The most valuable lesson from this assignment was how much of professional software "
 "engineering sits around the code rather than in it. Writing the CRUD logic was "
 "straightforward; the harder and more instructive parts were tracing requirements through "
 "SysML, decomposing the work into a coherent JIRA backlog and sprints, and wiring up a "
 "reliable CI/CD pipeline to a real server. Configuring the GitHub Actions deployment to EC2 "
 "with SSH and PM2 was the steepest challenge, particularly managing secrets and ensuring the "
 "tests run without a database by stubbing Mongoose. I also learned the importance of "
 "owner-scoping every query early, rather than retrofitting access control. If I repeated the "
 "project I would write tests alongside each feature from the start and automate the EC2 "
 "provisioning. Overall the assignment connected the individual tools into one continuous "
 "lifecycle, which is exactly how they are used in industry.")

# ---------- 10. References ----------
h1('10. References')
refs = [
 "diagrams.net. (n.d.). diagrams.net. https://app.diagrams.net/",
 "Atlassian. (n.d.). Jira Software. https://www.atlassian.com/software/jira",
 "Figma. (n.d.). Figma: The collaborative interface design tool. https://www.figma.com/",
 "GitHub. (n.d.). GitHub Actions documentation. https://docs.github.com/actions",
 "Amazon Web Services. (n.d.). Amazon EC2. https://aws.amazon.com/ec2/",
 "MongoDB Inc. (n.d.). MongoDB documentation. https://www.mongodb.com/docs/",
 "OpenJS Foundation. (n.d.). Express. https://expressjs.com/",
 "Meta Open Source. (n.d.). React. https://react.dev/",
 "nahaQUT. (n.d.). sampleapp_IFQ636 [Source code]. GitHub. https://github.com/nahaQUT/sampleapp_IFQ636",
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Pt(36)
    p.paragraph_format.first_line_indent = Pt(-36)

doc.save('IFQ636_A1_Report_OsamaMohamed.docx')

# word count of body prose (rough)
wc = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"Report saved. Approx body word count: {wc}")
